import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO
import re

def sanitize_string(input_string):
    # Remove hidden or control characters (ASCII < 32), except for \n, \r, and \t
    cleaned = re.sub(r'[^\x20-\x7E\n\r\t]', '', input_string)  # Keeps printable ASCII and \n, \r, \t
    
    # Remove any non-printable characters (ASCII 0-31 and 127)
    cleaned = re.sub(r'[\x00-\x1F\x7F]', '', cleaned)

    # Trim leading/trailing whitespace
    cleaned = cleaned.strip()

    return cleaned

# Function to get proposal content from OpenAI using raw API calls
def generate_proposal(requirements, proposer, proposed, timeline, cost):
    api_key = st.secrets["openai"]["api_key"]
    url = "https://api.openai.com/v1/chat/completions"

    prompt = "You are an expert proposal writer with extensive experience in crafting detailed and professional project proposals. Your task is to generate a comprehensive project proposal based on the following inputs. The proposal should include:\n\n1. **Executive Summary**: A brief overview of the project, its goals, and significance.\n2. **Project Requirements**: A detailed description of the project requirements, including specific needs and expectations.\n3. **Work Breakdown Structure**: A clear outline of tasks and deliverables associated with the project, categorized by phases or milestones.\n4. **Roles and Responsibilities**: Identification of key stakeholders and their respective roles in the project.\n5. **Timeline**: A realistic timeline for project completion, including major milestones.\n6. **Budget and Cost Estimates**: A breakdown of estimated costs associated with the project, including any resources required.\n7. **Risks and Mitigation Strategies**: Potential risks involved in the project and strategies for managing those risks. Generate a project proposal based on the following details:\nProject Requirements: "+requirements+"\nProposer Company: "+proposer+"\nProposed Company: "+proposed+"\nEstimated Timeline: "+timeline+"\nEstimated Cost: "+cost+"\nProvide a detailed work breakdown and requirement documentation in a professional format. Utilize athe above details to generate the actual proposal for me. Do not return the template as response instead return the actual proposal on the basis of the above details."

    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {api_key}'
    }

    payload = json.dumps({
        "model": "gpt-4o",
        "messages": [
            {
                "role": "system",
                "content": "You are an expert with many years of experience in creating proposals for various software development projects. Your expertise encompasses understanding client requirements, outlining project scope, defining deliverables, and estimating timelines and costs effectively. Provide comprehensive and well-structured proposals tailored to the specifics of each project."
            },
            {
                "role": "user",
                "content": sanitize_string(prompt)
            }
        ],
        "temperature": 0.5,
        "top_p": 0
    })

    response = requests.post(url, headers=headers, data=payload)
    
    if response.status_code == 200:
        return response.json()['choices'][0]['message']['content']
    else:
        return f"Error: {response.status_code}, {response.text}"

# Function to create a Word document
def create_word_doc(proposal_text):
    doc = Document()
    doc.add_heading('Project Proposal', 0)
    
    # Add content to the Word doc
    for line in proposal_text.split('\n'):
        if line.strip():  # Only add non-empty lines
            doc.add_paragraph(line)
    
    return doc

# Streamlit UI
st.title("Project Proposal Generator")

st.write("Enter the details to generate a project proposal:")

# Input fields for proposal details
requirements = st.text_area("Project Requirements", height=150, placeholder="Describe the project goals, features, and any specific functionalities needed.")
proposer = st.text_input("Proposer Company", placeholder="Enter the name of your company or organization.")
proposed = st.text_input("Proposed Company", placeholder="Enter the name of the company being proposed for the project.")
timeline = st.text_input("Estimated Timeline", placeholder="e.g., 3 months, Q1 2025, etc.")
cost = st.text_input("Estimated Cost", placeholder="e.g., $50,000, ₹20 Lakhs, etc.")

# Button to generate proposal
if st.button("Generate Proposal"):
    if requirements and proposer and proposed and timeline and cost:
        with st.spinner('Generating proposal...'):
            proposal_text = generate_proposal(requirements, proposer, proposed, timeline, cost)
        
        st.subheader("Generated Proposal")
        st.write(proposal_text)
        
        # Create a Word document
        doc = create_word_doc(proposal_text)
        
        # Prepare to download the document
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="Download Proposal as DOCX",
            data=buffer,
            file_name="project_proposal.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error("Please fill in all the fields to generate the proposal.")
